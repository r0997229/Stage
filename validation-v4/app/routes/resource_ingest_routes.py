#app\routes\resource_ingest_routes.py
from __future__ import annotations

import io
import json
from typing import Any, Dict, List, Optional
from openai import OpenAIError

from flask import Blueprint, current_app, redirect, render_template, request, session, url_for

from ..control.resources import doc_sections_template as dst_module
from ..control.resources.doc_sections_template import DOC_SECTIONS_TEMPLATE
from ..control.ai_call_param.template_param import (
    MODEL, EFFORT, VERBOSITY, SYSTEM, INSTRUCTIONS
)
from ..services.openai_api_call import call_openai
from ..services.resources.resource_persistence import persist_python_dict_constant

resource_ingest_bp = Blueprint("resource_ingest_bp", __name__)

ALLOWED_DOC_TYPES = {"VP", "RA", "VR"}
ALLOWED_EXTENSIONS = (".docx", ".pdf")


def _extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells: List[str] = []
            for cell in row.cells:
                cell_text = " ".join((cell.text or "").split())
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts).strip()


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber  # type: ignore

        pages: List[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                txt = (page.extract_text() or "").strip()
                if txt:
                    pages.append(txt)

        out = "\n\n".join(pages).strip()
        if out:
            return out
    except ImportError:
        pass

    try:
        import fitz  # type: ignore

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: List[str] = []
        for i in range(len(doc)):
            txt = (doc.load_page(i).get_text("text") or "").strip()
            if txt:
                pages.append(txt)
        return "\n\n".join(pages).strip()
    except ImportError as exc:
        raise RuntimeError(
            "PDF support missing. Install either:\n"
            "  pip install pdfplumber\n"
            "or:\n"
            "  pip install pymupdf"
        ) from exc


def _extract_text_any(file_storage) -> str:
    filename = (file_storage.filename or "").lower().strip()
    if not filename:
        raise ValueError("Missing filename.")
    if not filename.endswith(ALLOWED_EXTENSIONS):
        raise ValueError("Unsupported file type. Only .docx and .pdf are allowed.")

    file_bytes = file_storage.read()
    if not file_bytes:
        raise ValueError("Uploaded file is empty.")

    if filename.endswith(".docx"):
        return _extract_text_from_docx(file_bytes)
    return _extract_text_from_pdf(file_bytes)

@resource_ingest_bp.route("/template/ingest", methods=["GET"])
def ingest_template():
    selected = (request.args.get("doc_type") or "").strip()
    existing = DOC_SECTIONS_TEMPLATE.get(selected) if selected else None

    return render_template(
        "resources/template_ingest.html",
        title="Template ingest",
        toast_error=session.pop("toast_error", None),
        toast_success=session.pop("toast_success", None),
        selected_doc_type=selected,
        existing_template=existing,
    )


@resource_ingest_bp.route("/template/ingest", methods=["POST"])
def ingest_template_post():
    doc_type = (request.form.get("doc_type") or "").strip()
    overwrite = (request.form.get("overwrite") or "").lower() in ("1", "true", "yes", "on")
    if doc_type not in ALLOWED_DOC_TYPES:
        session["toast_error"] = "Invalid document type. Use VP, RA, or VR."
        return redirect(url_for("resource_ingest_bp.ingest_template"))

    if doc_type in DOC_SECTIONS_TEMPLATE and not overwrite:
        session["toast_error"] = f'Template "{doc_type}" already exists. Enable overwrite to replace it.'
        return redirect(url_for("resource_ingest_bp.ingest_template", doc_type=doc_type))

    f = request.files.get("file")
    if not f:
        session["toast_error"] = "Missing file."
        return redirect(url_for("resource_ingest_bp.ingest_template", doc_type=doc_type))

    try:
        template_text = _extract_text_any(f)
        if not template_text.strip():
            session["toast_error"] = "No text could be extracted from the document."
            return redirect(url_for("resource_ingest_bp.ingest_template", doc_type=doc_type))

        try:
            template_payload = call_openAI(template_text, MODEL, EFFORT, VERBOSITY, SYSTEM, INSTRUCTIONS)

            if template_payload is None:
                raise ValueError("Normalizer returned invalid JSON.")
        except OpenAIError:
            current_app.logger.exception("Template generation failed")
            session["toast_error"] = "AI Template generation failed."
            return redirect(url_for("resource_ingest_bp.ingest_template", doc_type=doc_type))

        DOC_SECTIONS_TEMPLATE[doc_type] = template_payload

        persist_python_dict_constant(
            module_file=dst_module.__file__,
            constant_name="DOC_SECTIONS_TEMPLATE",
            data=DOC_SECTIONS_TEMPLATE,
        )
        session["toast_success"] = f'Template "{doc_type}" saved ({"replaced" if overwrite else "added"}).'
        return redirect(url_for("resource_ingest_bp.ingest_template", doc_type=doc_type))

    except (ValueError, RuntimeError) as exc:
        # ValueError: input/validation issues
        # RuntimeError: missing PDF deps, etc.
        session["toast_error"] = str(exc)
        return redirect(url_for("resource_ingest_bp.ingest_template", doc_type=doc_type))

    except Exception:
        current_app.logger.exception("Template ingest failed")
        session["toast_error"] = "Internal error while ingesting template."
        return redirect(url_for("resource_ingest_bp.ingest_template", doc_type=doc_type))